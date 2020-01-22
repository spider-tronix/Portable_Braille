import re
import docx

"""
Starting code for post processing

Things which might need to be processed:

- figures
- tables 
- page numbers
- headings
- bullets
- Chapter, Lesson
"""


def postprocess(text):
    """Example function for figures only"""

    # Detect figures
    figures = re.finditer(r"(figure|fig|table)( ?)(.*)\.$", text, flags=re.IGNORECASE|re.DOTALL)
    # TODO : decide what to do with the detected figures

    # Detect bullets - all will be replaced by "*" bullet
    text = text.replace('/[#•»>‣◦⁃⁌⁍∙*]/gi', '*')

    return text

# to iterate over the headings and paragraphs
# def iter_headings(paragraphs):
#     for paragraph in paragraphs:
#         if paragraph.style.name.startswith('Heading'):
#             yield paragraph
#
# for heading in iter_headings(document.paragraphs):
#     print heading.text

# from docx import Document
#
# document = Document("North Sydney TE SP30062590-1 HVAC - Project Offer -  Rev1.docx")
#
# ind = [i for i, para in enumerate(document.paragraphs) if 'heading' in para.text]
# if ind:
#     for i, para in enumerate(document.paragraphs):
#         if i > ind[0]:
#              print(para.text)

# Get headings from the contents page or the index page of the book if we can

# regex's for headings

# Heading: (?m)^(\d+\.\d+\s[ \w,\-]+)\r?$
# Subheading: (?m)^(\d\.[\d.]+ ?[ \w]+) ?\r?$

# import pytesseract
#
# text = pytesseract.image_to_string("<-----------image path------------->")
# flag_newline = False
# flag_heading = False
# flag_new_head = False
# count = 0
# headings = {}
# for i in range(len(text)):
#     c = text[i]
#     if c == '\n':
#         if flag_newline:
#             count += 1
#         else:
#             count = 1
#             flag_newline = True
#         continue
#     elif count == 1:
#         if flag_newline:
#             flag_newline = False
#             count = 0
#         elif flag_heading:
#             count = 0
#             flag_newline = False
#             flag_heading = False
#     if count >= 2:
#         if c == ' ':
#             count += 1
#             continue
#         else:
#             if flag_heading:
#                 headings[i - count] = headings[None]
#                 flag_heading = False
#             else:
#                 headings[None] = i
#                 count = 0
#                 flag_heading = True
#                 flag_newline = False
#     elif count == 1 and flag_heading:
#         count = 0
#         flag_heading = False
#         flag_newline = False
#         flag_new_head = False


# heading_detect = r'[\n\s]*\n\n.*\n\n[\n\s]*'
# paragraphs = r'\s*.+\.\n'
# figures_and_tables = r"(figure|fig|table)( ?)(.*)\."


import cv2
import numpy as np

img = cv2.imread('E:\\My Desktop\\Spider\\portable-braille\\Images\\tesseract test\\IMG_20191005_095940.jpg', -1)

rgb_planes = cv2.split(img)

result_planes = []
result_norm_planes = []
for plane in rgb_planes:
    dilated_img = cv2.dilate(plane, np.ones((7,7), np.uint8))
    bg_img = cv2.medianBlur(dilated_img, 21)
    diff_img = 255 - cv2.absdiff(plane, bg_img)
    norm_img = cv2.normalize(diff_img,None, alpha=0, beta=255, norm_type=cv2.NORM_MINMAX, dtype=cv2.CV_8UC1)
    result_planes.append(diff_img)
    result_norm_planes.append(norm_img)

result = cv2.merge(result_planes)
result_norm = cv2.merge(result_norm_planes)

cv2.imwrite('shadows_out.png', result)
cv2.imwrite('shadows_out_norm.png', result_norm)

#------------pre process image---------------#

import pytesseract, numpy as np, cv2

#Gamma correction

import numpy as np
# import argparse
import cv2


def adjust_gamma(image, gamma=1.0):
    # build a lookup table mapping the pixel values [0, 255] to
    # their adjusted gamma values
    invGamma = 1.0 / gamma
    table = np.array([((i / 255.0) ** invGamma) * 255
                      for i in np.arange(0, 256)]).astype("uint8")

    # apply gamma correction using the lookup table
    return cv2.LUT(image, table)


# loop over various values of gamma
for gamma in np.arange(0.0, 3.5, 0.5):
    # ignore when gamma is 1 (there will be no change to the image)
    if gamma == 1:
        continue

    # apply gamma correction and show the images
    gamma = gamma if gamma > 0 else 0.1
    adjusted = adjust_gamma(original, gamma=gamma)
    cv2.putText(adjusted, "g={}".format(gamma), (10, 30),
                cv2.FONT_HERSHEY_SIMPLEX, 0.8, (0, 0, 255), 3)
    cv2.imshow("Images", np.hstack([original, adjusted]))
    cv2.waitKey(0)

# histogram equalization (darkens the image - not useful alone)
# might work in combination with the gamma correction

out = cv2.equalizeHist(original)
cv2.imwrite("equalizedHist.jpg", out)

# Normalization (might be the best solution)

out = cv2.normalize(org, None, alphs=20, beta=240, norm_type=cv2.NORM_MINMAX, dtype=cv2.CV_32F)
cv2.imwrite("norm1.jpg", out)

# table detection

import cv2
import pytesseract

img = cv2.imread('C:\\Users\\yp270\\Desktop\\Spider\\portable-braille\\Images\\tesseract test\\IMG_20191005_095839.jpg')

h, w, c = img.shape
boxes = pytesseract.image_to_boxes(img)
for b in boxes.splitlines():
    b = b.split(' ')
    img = cv2.rectangle(img, (int(b[1]), h - int(b[2])), (int(b[3]), h - int(b[4])), (0, 255, 0), 2)

cv2.imshow('img', img)
cv2.waitKey(0)

# convert curve to line through page transformation

import cv2
import numpy as np
import matplotlib.pyplot as plt

def page_transformation(image_path, starting_row=int(input("Enter the row from which the transformation has to start:")),
                        transformation_func=lambda x: x**0.2):
    image = cv2.imread(image_path, 0)
    output = np.zeros(image.shape, dtype=image.dtype)
    rows, cols = image.shape

    for i in range(cols):
        increment = int(transformation_func(i))
        col = image[starting_row-increment:, i]
        output[:, i] = cv2.resize(col, (1, output.shape[0]), interpolation=cv2.INTER_CUBIC)[:, 0]

    return output

# example for a square root curve
output = page_transformation("../Images/tesseract test/thresh.jpg")
cv2.imshow('transformed image', output)
cv2.imshow('original image', cv2.imread("../Images/tesseract test/thresh.jpg", 0))
cv2.waitKey(0)
cv2.destroyAllWindows()